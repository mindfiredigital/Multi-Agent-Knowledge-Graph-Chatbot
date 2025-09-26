"""Microsoft Word document processor."""

from io import BytesIO
from typing import List

from docx import Document

# from docx.document import Document as DocumentType

from .base import DocumentProcessor


class DocxProcessor(DocumentProcessor):
    """Processor for Microsoft Word documents."""

    def extract_text(self, content: bytes) -> List[str]:
        """Extract text from Word document content."""
        if not self.validate_content(content):
            raise ValueError("Invalid Word document content")

        try:
            doc = Document(BytesIO(content))
            paragraphs = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            # If no paragraphs found, return empty list
            return paragraphs

        except (IOError, ValueError, TypeError) as e:
            raise IOError(f"Failed to extract text from Word document: {e}") from e

    def get_supported_extensions(self) -> List[str]:
        """Get supported Word document extensions."""
        return [".docx"]

    def validate_content(self, content: bytes) -> bool:
        """Validate Word document content."""
        if not content:
            return False
        try:
            Document(BytesIO(content))
            return True
        except (IOError, ValueError, TypeError):
            return False
